from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "report-source.md"
OUTPUT = ROOT / "Company_OS_ANC_Presentation_Evidence_2026.docx"

NAVY = "18324A"
BLUE = "2366A8"
PALE = "EAF2F8"
PALE_ALT = "F6F8FA"
WHITE = "FFFFFF"
GRAY = "5B6573"
LIGHT_BORDER = "D9D9D9"
CJK_FONT = "Noto Sans CJK SC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LIGHT_BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


INLINE_RE = re.compile(r"(\[([^\]]+)\]\((https?://[^)]+)\)|\*\*([^*]+)\*\*)")


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        if match.group(2) and match.group(3):
            add_hyperlink(paragraph, match.group(2), match.group(3))
        else:
            run = paragraph.add_run(match.group(4))
            run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Company OS · Evidence Brief  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CJK_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string("27313C")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color, before, after in (
        ("Title", 31, NAVY, 0, 14),
        ("Subtitle", 14, GRAY, 0, 12),
        ("Heading 1", 20, NAVY, 12, 7),
        ("Heading 2", 14, BLUE, 9, 5),
        ("Heading 3", 11.5, NAVY, 7, 3),
    ):
        style = styles[name]
        style.font.name = CJK_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    # Remove the decorative border inherited by some built-in Title styles.
    title_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    for name in ("Evidence", "Stage line"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    evidence = styles["Evidence"]
    evidence.base_style = styles["Normal"]
    evidence.font.size = Pt(8.5)
    evidence.font.color.rgb = RGBColor.from_string(GRAY)
    evidence.paragraph_format.left_indent = Inches(0.18)
    evidence.paragraph_format.space_after = Pt(5)
    stage = styles["Stage line"]
    stage.base_style = styles["Normal"]
    stage.font.size = Pt(11)
    stage.font.bold = True
    stage.font.color.rgb = RGBColor.from_string(NAVY)

    header = section.header.paragraphs[0]
    header.text = "COMPANY OS / ANC  ·  PRESENTATION EVIDENCE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Inches(0.8)
    r = p.add_run("COMPANY OS / ANC")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph(style="Title")
    title.add_run("演示证据报告")
    title_ppr = title._p.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run("为什么企业需要一个跨 Agent、跨模型、跨部署方式的公司控制平面")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Inches(0.28)
    intro.paragraph_format.space_after = Inches(0.35)
    run = intro.add_run(
        "市场采用 · ROI · Agent 治理 · 身份与权限 · 互操作标准 · 安全风险 · 全球法规 · 中国市场 · 产品验证"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    widths = (Inches(1.55), Inches(5.35))
    rows = [
        ("核心判断", "企业缺的不是另一个 Agent，而是管理所有 Agent 的组织、权限、审批、证据与责任基础设施。"),
        ("目标听众", "企业客户、合作伙伴、投资人、内部决策者"),
        ("研究范围", "全球 + 中国；重点覆盖 2023–2026 年的官方、标准、研究和市场证据"),
        ("版本日期", "2026-09-04"),
    ]
    for i, (key, value) in enumerate(rows):
        for j, cell in enumerate(table.rows[i].cells):
            cell.width = widths[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_shading(cell, PALE if j == 0 else WHITE)
        table.cell(i, 0).paragraphs[0].add_run(key).bold = True
        table.cell(i, 1).paragraphs[0].add_run(value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Inches(0.55)
    r = p.add_run("研究纪律")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.add_run("法规/标准和独立研究用于核心结论；厂商调查用于趋势与品类验证；Company OS 自有证据只证明产品实现，不替代第三方客户案例。")
    p = doc.add_paragraph()
    p.add_run("本报告中的全部蓝色来源名称均为可点击链接。")
    doc.add_page_break()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for i, row in enumerate(rows):
        for j in range(column_count):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if i == 0:
                set_cell_shading(cell, NAVY)
            elif i % 2 == 0:
                set_cell_shading(cell, PALE_ALT)
            text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(8)
                if i == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def should_page_break_before(title: str) -> bool:
    return title in {
        "核心证据卡",
        "Company OS 与证据的对应关系",
        "竞品与品类信号",
        "反对意见与证据回答",
        "完整链接库（按可信度与用途整理）",
        "现场演示建议",
        "不应过度声称的内容",
    }


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    # Skip the source title and metadata already represented on the cover.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 一句话结论"))
    i = start
    slide_counter = 0
    in_slide_story = False
    current_section = ""
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue
        if stripped.startswith("## "):
            title = stripped[3:]
            current_section = title
            if should_page_break_before(title):
                doc.add_page_break()
            doc.add_heading(title, level=1)
            in_slide_story = title == "10 页演示主线"
            i += 1
            continue
        if stripped.startswith("### "):
            title = stripped[4:]
            if in_slide_story:
                slide_counter += 1
            doc.add_heading(title, level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            if current_section == "现场演示建议":
                p = doc.add_paragraph()
                prefix = p.add_run(f"{numbered.group(1)}. ")
                prefix.bold = True
                add_inline(p, numbered.group(2))
            else:
                p = doc.add_paragraph(style="List Number")
                add_inline(p, numbered.group(2))
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            i += 1
            continue
        p_style = None
        if stripped.startswith("**台上说法：**"):
            p_style = "Stage line"
        elif stripped.startswith("**证据：**"):
            p_style = "Evidence"
        p = doc.add_paragraph(style=p_style)
        add_inline(p, stripped)
        i += 1


def set_document_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Company OS / ANC 演示证据报告"
    props.subject = "AI Agent 企业控制平面的市场、治理、安全、法规和产品证据"
    props.author = "Company OS Research"
    props.keywords = "Company OS, ANC, AI agents, governance, identity, evidence, control plane"
    props.comments = "Research snapshot as of 2026-09-04; verify dynamic web sources before external publication."


def force_unicode_font(doc: Document) -> None:
    """Make the rendered QA copy deterministic on macOS LibreOffice."""
    roots = [doc._element]
    roots.extend(section.header._element for section in doc.sections)
    roots.extend(section.footer._element for section in doc.sections)
    for root in roots:
        for run in root.iter(qn("w:r")):
            r_pr = run.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                run.insert(0, r_pr)
            fonts = r_pr.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                r_pr.insert(0, fonts)
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{attr}"), CJK_FONT)


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    set_document_metadata(doc)
    add_cover(doc)
    add_markdown_body(doc, markdown)

    # Prevent widows/orphans where Word supports it.
    for paragraph in doc.paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)
    force_unicode_font(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
